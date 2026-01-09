import streamlit as st
import pandas as pd
import numpy as np
from datetime import datetime
import io
import streamlit.components.v1 as components
from docx import Document
from docx.shared import Inches

# 嘗試匯入 Firebase 套件
try:
    import firebase_admin
    from firebase_admin import credentials, firestore, auth, initialize_app, get_app
    HAS_FIREBASE = True
except ImportError:
    HAS_FIREBASE = False

# 頁面配置
st.set_page_config(page_title="正覺壁球管理系統", layout="wide", initial_sidebar_state="expanded")

# --- 1. Firebase 初始化 ---
def init_firebase():
    if not HAS_FIREBASE: return None
    if 'firebase_initialized' not in st.session_state:
        try:
            try:
                app = get_app()
            except ValueError:
                if "firebase_config" in st.secrets:
                    key_dict = dict(st.secrets["firebase_config"])
                    if "private_key" in key_dict:
                        key_dict["private_key"] = key_dict["private_key"].replace("\\n", "\n")
                    cred = credentials.Certificate(key_dict)
                    app = initialize_app(cred)
                else: return None
            st.session_state.db = firestore.client()
            st.session_state.firebase_initialized = True
        except Exception as e:
            st.error(f"Firebase 初始化失敗: {e}")
            return None
    return st.session_state.get('db')

db = init_firebase()
app_id = "squash-management-v1"

# --- 2. 數據存取與「預設數據」邏輯 ---
def get_default_df(name):
    """還原原稿中的所有初始預設數據，確保代碼豐滿度與功能性"""
    if name == 'schedules':
        return pd.DataFrame([
            {"日期": "2024-03-20", "時間": "16:00-17:30", "班別": "校隊班", "地點": "太和體育館", "教練": "陳教練"},
            {"日期": "2024-03-21", "時間": "16:00-17:30", "班別": "精英班", "地點": "太和體育館", "教練": "林教練"},
            {"日期": "2024-03-22", "時間": "15:30-16:30", "班別": "興趣班", "地點": "學校操場", "教練": "黃教練"}
        ])
    elif name == 'rankings':
        return pd.DataFrame([
            {"姓名": "張小明", "班級": "5A", "積分": 1200, "章別": "金章"},
            {"姓名": "李華", "班級": "4C", "積分": 1150, "章別": "銀章"},
            {"姓名": "王小城", "班級": "6B", "積分": 980, "章別": "銀章"},
            {"姓名": "趙大衛", "班級": "3A", "積分": 850, "章別": "銅章"}
        ])
    elif name == 'awards':
        return pd.DataFrame([
            {"獲獎日期": "2023-12-01", "學生姓名": "張小明", "比賽名稱": "全港青少年壁球錦標賽", "獎項": "男子U12季軍"},
            {"獲獎日期": "2024-01-15", "學生姓名": "李華", "比賽名稱": "新界區分齡賽", "獎項": "殿軍"}
        ])
    elif name == 'attendance':
        return pd.DataFrame(columns=["日期", "班別", "出席名單", "人數", "備註"])
    elif name == 'news':
        return pd.DataFrame([
            {"日期": "2024-03-01", "標題": "復活節訓練安排", "內容": "復活節假期期間，所有校隊訓練暫停一次，改為自主練習。"},
            {"日期": "2024-02-28", "標題": "新隊服發放", "內容": "請已訂購隊服的隊員於週三訓練後到更衣室領取。"}
        ])
    elif name == 'tournaments':
        return pd.DataFrame([
            {"截止日期": "2024-04-10", "賽事名稱": "全港校際壁球挑戰賽", "報名連結": "https://example.com/reg1", "狀態": "報名中"},
            {"截止日期": "2024-05-01", "賽事名稱": "暑期青少年盃", "報名連結": "", "狀態": "即將開放"}
        ])
    return pd.DataFrame()

def fetch_data(collection_name, default_cols):
    cache_key = f"cache_{collection_name}"
    if db:
        try:
            coll_ref = db.collection('artifacts').document(app_id).collection('public').document('data').collection(collection_name)
            docs = coll_ref.stream()
            items = [doc.to_dict() for doc in docs]
            if items:
                df = pd.DataFrame(items)
                st.session_state[cache_key] = df
                return df
        except: pass
    
    if cache_key not in st.session_state:
        st.session_state[cache_key] = get_default_df(collection_name)
    return st.session_state[cache_key]

def save_data(collection_name, df):
    if df is None: return
    st.session_state[f"cache_{collection_name}"] = df
    if db:
        try:
            coll_ref = db.collection('artifacts').document(app_id).collection('public').document('data').collection(collection_name)
            old_docs = coll_ref.stream()
            for d in old_docs: d.reference.delete()
            for i, row in df.iterrows():
                data_item = {k: (v if pd.notna(v) else None) for k, v in row.to_dict().items()}
                coll_ref.document(f"doc_{i}").set(data_item)
            st.toast(f"✅ {collection_name} 雲端同步成功")
        except Exception as e: st.error(f"寫入失敗: {e}")

# --- 3. 樣式與側邊欄 ---
st.markdown("""
<style>
    .main { background: linear-gradient(135deg, #f5f7fa 0%, #c3cfe2 100%); }
    .stApp { background-attachment: fixed; }
    .reportview-container .main .block-container { padding-top: 2rem; }
    div[data-testid="stMetricValue"] { color: #2563eb; font-weight: 800; }
    .stSidebar { background-color: #ffffff !important; box-shadow: 2px 0 10px rgba(0,0,0,0.05); }
    .stButton>button { border-radius: 20px; transition: all 0.3s; }
    .stButton>button:hover { transform: translateY(-2px); box-shadow: 0 4px 12px rgba(0,0,0,0.1); }
    .announcement-card { 
        background: white; padding: 25px; border-radius: 15px; border-left: 6px solid #2563eb;
        margin-bottom: 20px; box-shadow: 0 4px 6px rgba(0,0,0,0.05);
    }
    .rank-card {
        background: white; padding: 20px; border-radius: 15px; text-align: center;
        border: 1px solid #e2e8f0; transition: all 0.3s;
    }
    .rank-card:hover { border-color: #2563eb; transform: scale(1.02); }
</style>
""", unsafe_allow_html=True)

if 'auth' not in st.session_state:
    st.session_state.auth = {"logged_in": False, "user_id": "", "role": "guest"}

st.sidebar.markdown("## 🏸 正覺壁球管理系統")
st.sidebar.image("https://cdn-icons-png.flaticon.com/512/3253/3253041.png", width=100)

if not st.session_state.auth["logged_in"]:
    st.sidebar.subheader("🔒 系統登入")
    login_type = st.sidebar.radio("身份選擇", ["學生/家長", "教練/管理員"])
    if login_type == "教練/管理員":
        pwd = st.sidebar.text_input("後台管理密碼", type="password")
        if st.sidebar.button("進入管理後台"):
            if pwd == "8888":
                st.session_state.auth = {"logged_in": True, "user_id": "ADMIN", "role": "admin"}
                st.rerun()
            else: st.sidebar.error("密碼錯誤")
    else:
        s_class = st.sidebar.text_input("班別 (e.g., 4A)")
        s_no = st.sidebar.text_input("學號 (e.g., 01)")
        if st.sidebar.button("查詢登入"):
            if s_class and s_no:
                st.session_state.auth = {"logged_in": True, "user_id": f"{s_class}_{s_no}", "role": "student"}
                st.rerun()
    st.info("💡 學生登入可查看個人進度與公告")
    st.stop()

# --- 4. 導航選單 ---
menu = [
    "📅 訓練日程表", "🏆 隊員排行榜", "🤖 AI 動作深度分析", 
    "📝 考勤點名中心", "🎖️ 學生得獎紀錄", "📢 隊內最新公告", 
    "⚡ 比賽報名系統", "📊 營運預算核算"
]
choice = st.sidebar.radio("📌 主選單", menu)

# --- 5. 功能模組 ---

# A. 訓練日程表 (完全還原複雜邏輯)
if choice == "📅 訓練日程表":
    st.title("📅 壁球訓練班日程安排")
    sched_df = fetch_data('schedules', [])
    
    col1, col2 = st.columns([3, 1])
    with col1:
        st.markdown("### 🗓️ 本期課表")
        f_class = st.multiselect("篩選班別", options=list(sched_df["班別"].unique()))
        display_df = sched_df.copy()
        if f_class: display_df = display_df[display_df["班別"].isin(f_class)]
        st.dataframe(display_df, use_container_width=True, hide_index=True)
        
    with col2:
        st.markdown("### 🔧 數據操作")
        if st.session_state.auth["role"] == "admin":
            if st.button("➕ 新增訓練堂數"):
                new_row = {"日期": str(datetime.now().date()), "時間": "16:00-17:30", "班別": "未定", "地點": "太和", "教練": "陳教練"}
                sched_df = pd.concat([sched_df, pd.DataFrame([new_row])], ignore_index=True)
                save_data('schedules', sched_df)
                st.rerun()
            
            output = io.BytesIO()
            with pd.ExcelWriter(output, engine='openpyxl') as writer:
                display_df.to_excel(writer, index=False)
            st.download_button("📥 匯出 Excel 課表", output.getvalue(), "schedule.xlsx")

    if st.session_state.auth["role"] == "admin":
        with st.expander("📝 批次編輯模式"):
            edited_df = st.data_editor(sched_df, num_rows="dynamic", use_container_width=True)
            if st.button("💾 儲存並同步雲端"):
                save_data('schedules', edited_df)
                st.rerun()

# B. 排行榜 (精美卡片還原)
elif choice == "🏆 隊員排行榜":
    st.title("🏆 正覺壁球隊榮譽排行榜")
    rank_df = fetch_data('rankings', [])
    rank_df["積分"] = pd.to_numeric(rank_df["積分"], errors='coerce').fillna(0)
    sorted_rank = rank_df.sort_values("積分", ascending=False).reset_index(drop=True)

    # 頂部三甲
    top_cols = st.columns(3)
    medals = ["🥇 冠軍", "🥈 亞軍", "🥉 季軍"]
    colors = ["#ffd700", "#c0c0c0", "#cd7f32"]
    for i in range(min(3, len(sorted_rank))):
        with top_cols[i]:
            row = sorted_rank.iloc[i]
            st.markdown(f"""
            <div class="rank-card" style="border-top: 5px solid {colors[i]};">
                <p style="font-size: 1.2rem; color: #64748b;">{medals[i]}</p>
                <h2 style="margin: 10px 0;">{row['姓名']}</h2>
                <h3 style="color: #2563eb;">{int(row['積分'])} pts</h3>
                <span style="background: #f1f5f9; padding: 4px 12px; border-radius: 10px;">{row['班級']} | {row['章別']}</span>
            </div>
            """, unsafe_allow_html=True)
    
    st.markdown("### 📊 全員積分明細")
    st.table(sorted_rank)

# C. AI 動作分析 (新增)
elif choice == "🤖 AI 動作深度分析":
    st.title("🤖 AI 動作深度分析系統")
    st.info("此系統運用 MediaPipe 視覺技術，協助分析揮拍角度。")
    ai_script = """
    <div style="background:#fff; padding:20px; border-radius:15px; box-shadow: 0 4px 12px rgba(0,0,0,0.1);">
        <script src="https://cdn.jsdelivr.net/npm/@mediapipe/pose/pose.js"></script>
        <script src="https://cdn.jsdelivr.net/npm/@mediapipe/drawing_utils/drawing_utils.js"></script>
        <input type="file" id="video-in" accept="video/*" style="margin-bottom:10px; width:100%; padding:10px; border:1px dashed #2563eb;">
        <div style="position:relative;">
            <video id="v-src" controls style="width:100%; border-radius:10px;"></video>
            <canvas id="v-canvas" style="position:absolute; top:0; left:0; width:100%; height:100%; pointer-events:none;"></canvas>
        </div>
        <div id="status" style="margin-top:15px; padding:15px; border-radius:10px; background:#f8fafc; font-weight:bold; color:#2563eb; text-align:center;">
            手肘角度：<span id="ang-val">0</span>°
        </div>
    </div>
    <script>
        const v = document.getElementById('v-src');
        const c = document.getElementById('v-canvas');
        const ctx = c.getContext('2d');
        const angVal = document.getElementById('ang-val');
        const pose = new Pose({locateFile: (f) => `https://cdn.jsdelivr.net/npm/@mediapipe/pose/${f}`});
        pose.setOptions({modelComplexity: 1, smoothLandmarks: true, minDetectionConfidence: 0.5});
        pose.onResults((res) => {
            if (!res.poseLandmarks) return;
            c.width = v.clientWidth; c.height = v.clientHeight;
            ctx.clearRect(0,0,c.width,c.height);
            drawConnectors(ctx, res.poseLandmarks, POSE_CONNECTIONS, {color: '#00FF00', lineWidth: 3});
            const s = res.poseLandmarks[12], e = res.poseLandmarks[14], w = res.poseLandmarks[16];
            if (s && e && w) {
                const r = Math.atan2(w.y-e.y, w.x-e.x) - Math.atan2(s.y-e.y, s.x-e.x);
                let d = Math.abs(r * 180 / Math.PI);
                if (d > 180) d = 360 - d;
                angVal.innerText = d.toFixed(1);
            }
        });
        document.getElementById('video-in').onchange = (e) => {
            const f = e.target.files[0];
            if (f) { v.src = URL.createObjectURL(f); v.play(); }
        };
        async function run() {
            if (!v.paused && !v.ended) { await pose.send({image: v}); }
            requestAnimationFrame(run);
        }
        v.onplay = run;
    </script>
    """
    components.html(ai_script, height=700)

# D. 考勤點名 (完全還原表單與報表)
elif choice == "📝 考勤點名中心":
    st.title("📝 考勤與訓練記錄")
    att_df = fetch_data('attendance', [])
    
    if st.session_state.auth["role"] == "admin":
        with st.form("att_form", clear_on_submit=True):
            st.subheader("➕ 記錄今日訓練")
            c1, c2, c3 = st.columns(3)
            date = c1.date_input("訓練日期", datetime.now())
            cls = c2.selectbox("訓練班別", ["校隊班", "精英班", "中級班", "初級班", "興趣班"])
            note = c3.text_input("課堂備註 (如: 學習長球)")
            names = st.text_area("出席學生名單 (請以空格或逗號分隔)")
            submit = st.form_submit_button("✅ 儲存考勤")
            
            if submit:
                clean_names = [n.strip() for n in names.replace(',', ' ').split() if n.strip()]
                new_att = {"日期": str(date), "班別": cls, "出席名單": ", ".join(clean_names), "人數": len(clean_names), "備註": note}
                att_df = pd.concat([att_df, pd.DataFrame([new_att])], ignore_index=True)
                save_data('attendance', att_df)
                st.success(f"已記錄 {len(clean_names)} 位隊員出席")
                st.rerun()

    st.markdown("### 歷史考勤表")
    st.dataframe(att_df.sort_values("日期", ascending=False), use_container_width=True, hide_index=True)
    
    # 匯出 Word 報告邏輯
    if not att_df.empty:
        if st.button("📄 生成本月考勤報告 (Word)"):
            doc = Document()
            doc.add_heading('正覺壁球隊 - 月度考勤報告', 0)
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '日期'
            hdr_cells[1].text = '班別'
            hdr_cells[2].text = '人數'
            hdr_cells[3].text = '出席名單'
            for _, row in att_df.iterrows():
                row_cells = table.add_row().cells
                row_cells[0].text = str(row['日期'])
                row_cells[1].text = row['班別']
                row_cells[2].text = str(row['人數'])
                row_cells[3].text = row['出席名單']
            doc_io = io.BytesIO()
            doc.save(doc_io)
            st.download_button("📥 下載 Word 報告", doc_io.getvalue(), "attendance_report.docx")

# E. 得獎紀錄
elif choice == "🎖️ 學生得獎紀錄":
    st.title("🎖️ 榮譽榜")
    award_df = fetch_data('awards', [])
    for _, row in award_df.iterrows():
        st.markdown(f"""
        <div class="announcement-card" style="border-left-color: #ffd700;">
            <p style="color: #64748b; font-size: 0.8rem;">{row['獲獎日期']}</p>
            <h3 style="margin: 5px 0;">{row['比賽名稱']}</h3>
            <p>🏆 <b>{row['獎項']}</b> — {row['學生姓名']}</p>
        </div>
        """, unsafe_allow_html=True)
    if st.session_state.auth["role"] == "admin":
        with st.expander("編輯獲獎數據"):
            save_data('awards', st.data_editor(award_df, num_rows="dynamic"))

# F. 公告欄 (還原樣式)
elif choice == "📢 隊內最新公告":
    st.title("📢 最新消息")
    news_df = fetch_data('news', [])
    for _, row in news_df.iloc[::-1].iterrows():
        st.markdown(f"""
        <div class="announcement-card">
            <small style="color: #2563eb; font-weight: bold;">{row['日期']}</small>
            <h3 style="margin-top: 5px;">{row['標題']}</h3>
            <p style="color: #475569; line-height: 1.6;">{row['內容']}</p>
        </div>
        """, unsafe_allow_html=True)

# G. 比賽報名
elif choice == "⚡ 比賽報名系統":
    st.title("⚡ 賽事報名直通車")
    tour_df = fetch_data('tournaments', [])
    for _, row in tour_df.iterrows():
        c1, c2 = st.columns([4, 1])
        with c1:
            st.markdown(f"**{row['賽事名稱']}**")
            st.caption(f"截止日期: {row['截止日期']} | 狀態: {row['狀態']}")
        with c2:
            if row['報名連結']: st.link_button("前往報名", row['報名連結'])
            else: st.button("尚未開放", disabled=True)
        st.divider()

# H. 營運預算 (完全還原原稿的所有費率與公式)
elif choice == "📊 營運預算核算":
    st.title("📊 隊伍財務預算與估算")
    if st.session_state.auth["role"] != "admin":
        st.warning("🔒 此功能僅限教練/管理員訪問。")
        st.stop()
        
    c1, c2 = st.columns(2)
    with c1:
        st.subheader("🛠️ 支出預測 (教練費)")
        n1 = st.number_input("校隊班 (2750/次) - 堂數", value=11)
        n2 = st.number_input("精英/中/初級班 (1350/次) - 堂數", value=30)
        n3 = st.number_input("興趣班 (1200/次) - 堂數", value=24)
        total_exp = (n1 * 2750) + (n2 * 1350) + (n3 * 1200)
        
    with c2:
        st.subheader("💰 收入預測 (學費)")
        std_count = st.number_input("預計收生總人數", value=60)
        fee = st.number_input("平均每人學費 ($)", value=800)
        total_rev = std_count * fee
        
    st.divider()
    m1, m2, m3 = st.columns(3)
    m1.metric("預計總收入", f"${total_rev:,}")
    m2.metric("預計開班總支出", f"${total_exp:,}")
    m3.metric("預計利潤/餘額", f"${total_rev - total_exp:,}", delta=f"{total_rev - total_exp}")
    
    st.bar_chart(pd.DataFrame({"金額": [total_rev, total_exp]}, index=["收入", "支出"]))

# 頁腳
st.sidebar.markdown("---")
st.sidebar.caption("正覺壁球管理系統 v1.2.0")
st.sidebar.caption("© 2024 壁球隊行政組")
